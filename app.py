import os
from dotenv import load_dotenv
from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
from langdetect import detect
from google.cloud import translate_v2 as translate
from docx import Document
import pdfkit
from PyPDF2 import PdfReader
import pytesseract
import chardet

app = Flask(__name__)
load_dotenv()

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

# Load credentials from environment variables
credentials_path = os.getenv('GOOGLE_APPLICATION_CREDENTIALS')
if credentials_path is None:
    raise ValueError("GOOGLE_APPLICATION_CREDENTIALS environment variable is not set.")

# Initialize the Google Translator client
translate_client = translate.Client.from_service_account_json(credentials_path)

def save_file(file):
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    return file_path

def detect_language(file_path):
    if file_path.endswith('.docx'):
        text = extract_text_from_docx(file_path)
    else:
        # Detect encoding for text files
        with open(file_path, 'rb') as file:
            raw_data = file.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding']

        # Read the file with the detected encoding
        with open(file_path, 'r', encoding=encoding) as file:
            text = file.read()

    return detect(text)

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return '\n'.join(text)

def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def translate_text(text, target_language):
    max_chars_per_request = 100000
    translated_text = ""

    for i in range(0, len(text), max_chars_per_request):
        chunk = text[i:i + max_chars_per_request]
        translated = translate_client.translate(chunk, target_language)
        translated_text += translated['translatedText'] + " "

    return translated_text.strip()

def save_as_docx(text, filename):
    doc = Document()
    doc.add_paragraph(text)
    doc_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{filename}.docx")
    doc.save(doc_path)
    return doc_path

def save_as_txt(text, filename):
    txt_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{filename}.txt")
    with open(txt_path, 'w', encoding='utf-8') as file:
        file.write(text)
    return txt_path

def save_as_pdf(html, filename):
    pdf_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{filename}.pdf")
    pdfkit.from_string(html, pdf_path)
    return pdf_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    files = request.files.getlist('document')
    target_language = request.form['language']
    output_format = request.form['format']
    
    processed_files = []

    if not os.path.exists(app.config['OUTPUT_FOLDER']):
        os.makedirs(app.config['OUTPUT_FOLDER'])

    for file in files:
        file_path = save_file(file)

        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        elif file_path.endswith('.docx'):
            content = extract_text_from_docx(file_path)
        elif file_path.endswith('.pdf'):
            content = extract_text_from_pdf(file_path)
        else:
            content = pytesseract.image_to_string(file_path)

        translated_text = translate_text(content, target_language)

        filename = os.path.splitext(os.path.basename(file_path))[0]
        if output_format == 'doc':
            output_file_path = save_as_docx(translated_text, filename)
        elif output_format == 'txt':
            output_file_path = save_as_txt(translated_text, filename)
        elif output_format == 'pdf':
            output_file_path = save_as_pdf(translated_text, filename)
        
        processed_files.append(output_file_path)
    
    return send_file(processed_files[0], as_attachment=True)

# Vercel handler
def handler(event, context):
    return app(event, context)
